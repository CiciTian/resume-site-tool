"""Résumé file → raw text extraction.

Supports PDF, DOCX, TXT and Markdown. The output is plain text that gets handed
to the LLM extractor; no structure is inferred here.
"""

from __future__ import annotations

import io
from pathlib import Path


class ParseError(Exception):
    """Raised when a file cannot be turned into usable text."""


def _extract_pdf(data: bytes) -> str:
    import pdfplumber

    pages: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    import docx  # python-docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _extract_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def extract_text(filename: str, data: bytes) -> str:
    """Dispatch on file extension and return extracted text.

    Raises ParseError if the type is unsupported or the result is empty
    (e.g. a scanned PDF with no text layer — that needs OCR, out of scope).
    """
    suffix = Path(filename).suffix.lower()

    try:
        if suffix == ".pdf":
            text = _extract_pdf(data)
        elif suffix in (".docx",):
            text = _extract_docx(data)
        elif suffix in (".txt", ".md", ".markdown"):
            text = _extract_text(data)
        else:
            raise ParseError(
                f"Unsupported file type '{suffix}'. "
                "Supported: .pdf, .docx, .txt, .md"
            )
    except ParseError:
        raise
    except Exception as exc:  # noqa: BLE001 — surface a clean message to the API
        raise ParseError(f"Failed to read {suffix or 'file'}: {exc}") from exc

    text = text.strip()
    if not text:
        raise ParseError(
            "No text could be extracted. If this is a scanned PDF it needs OCR, "
            "which this tool does not do yet."
        )
    return text
